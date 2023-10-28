import mysql.connector
import argparse
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
import smtplib
import os

def ayuda():
    print("Para conocer la funcionalidad del cada control debes colocar por la linea de comandos:")
    print("py. proyecto.py --inventario-ayuda")
    print("py. proyecto.py --clientes-ayuda")
    print("py. proyecto.py --ventas-ayuda")
    print("Para generar y enviar los reportes de ventas deberás escribir el codigo de la siguiente manera:")
    print("py proyecto.py --generar-informe-ventas-cliente cliente_id destinatario")
    print("py proyecto.py --informe-ventas-producto producto_id destinatario")
    print("Para abrir el programa desde un menú en consola debes ingresar por la linea de comandos")
    print("py. proyecto.py --menu-interactivo")

def inventario_ayuda():
    print("Para poder utilizar el programa de Control de Inventario debes ingresar por la linea de comandos los datos separados por espacios, por ejemplo:")
    print("py proyecto.py --listar")
    print("py proyecto.py --crear codigo nombre existencia 'proveedor' precio")
    print("py proyecto.py --actualizar codigo nuevo_nombre nueva_existencia 'nuevo_proveedor' nuevo_precio" )
    print("py proyecto.py --editar-existencia codigo nueva_existencia")
    print("py proyecto.py --eliminar codigo")

# Función para conectar a la base de datos MySQL
def connect_to_database():
    try:
        connection = mysql.connector.connect(
            host="localhost",
            user="root",
            password="123456789",
            database="ventas"
        )
        return connection
    except mysql.connector.Error as error:
        print("Ocurrió un error al conectar a la base de datos:", error)
        return None

#CONTROL INVENTARIO 
def listar_productos(connection):
    try:
        cursor = connection.cursor()
        cursor.execute("SELECT codigo, nombre, existencia, proveedor, precio FROM inventario")
        productos = cursor.fetchall()
        for producto in productos:
            print("Código: {}, Nombre: {}, Existencia: {}, Proveedor: {}, Precio: {}".format(*producto))
    except mysql.connector.Error as error:
        print("Error al listar productos:", error)

def crear_producto(connection, codigo, nombre, existencia, proveedor, precio):
    try:
        cursor = connection.cursor()
        insert_query = "INSERT INTO inventario (codigo, nombre, existencia, proveedor, precio) VALUES (%s, %s, %s, %s, %s)"
        data = (codigo, nombre, existencia, proveedor, precio)
        cursor.execute(insert_query, data)
        connection.commit()
        print("Producto creado con éxito.")
    except mysql.connector.Error as error:
        print("Error al crear producto:", error)

# Función para actualizar un producto en la base de datos MySQL
def actualizar_producto(connection, codigo, nuevo_nombre, nueva_existencia, nuevo_proveedor, nuevo_precio):
    try:
        cursor = connection.cursor()
        update_query = "UPDATE inventario SET nombre=%s, existencia=%s, proveedor=%s, precio=%s WHERE codigo=%s"
        data = (nuevo_nombre, nueva_existencia, nuevo_proveedor, nuevo_precio, codigo)
        cursor.execute(update_query, data)
        connection.commit()
        print("Producto actualizado con éxito.")
    except mysql.connector.Error as error:
        print("Error al actualizar producto:", error)

def editar_existencia(connection, codigo, nueva_existencia):
    try:
        cursor = connection.cursor()
        update_query = "UPDATE inventario SET existencia=%s WHERE codigo=%s"
        data = (nueva_existencia, codigo)
        cursor.execute(update_query, data)
        connection.commit()
        print("Existencia actualizada con éxito.")
    except mysql.connector.Error as error:
        print("Error al editar existencia:", error)

def eliminar_producto(connection, codigo):
    try:
        cursor = connection.cursor()
        delete_query = "DELETE FROM inventario WHERE codigo=%s"
        data = (codigo,)
        cursor.execute(delete_query, data)
        connection.commit()
        print("Producto eliminado con éxito.")
    except mysql.connector.Error as error:
        print("Error al eliminar producto:", error)

def clientes_ayuda():
    print("Para poder utilizar el programa de Control de Clientes debe ingresar por la linea de comandos los datos separados por espacios, por ejemplo:")
    print("py proyecto.py --clientes-listar")
    print("py proyecto.py --clientes-crear codigo 'nombre' 'direccion' ")
    print("py proyecto.py --clientes-actualizar codigo nuevo_nombre nueva_direccion" )
    print("py proyecto.py --clientes-eliminar codigo")


#CONTROL CLIENTES
def listar_clientes(connection):
    try:
        cursor = connection.cursor()
        cursor.execute("SELECT codigo, nombre, direccion FROM clientes")
        clientes = cursor.fetchall()
        for clientes in clientes:
            print("Código: {}, Nombre: {}, Direccioón: {}".format(*clientes))
    except mysql.connector.Error as error:
        print("Error al listar clientes:", error)

def crear_clientes(connection, codigo, nombre, direccion):
    try:
        cursor = connection.cursor()
        insert_query = "INSERT INTO clientes (codigo, nombre, direccion) VALUES (%s, %s, %s)"
        data = (codigo, nombre, direccion)
        cursor.execute(insert_query, data)
        connection.commit()
        print("Cliente creado con éxito.")
    except mysql.connector.Error as error:
        print("Error al crear Cliente:", error)

# Función para actualizar un cliente en la base de datos MySQL
def actualizar_clientes(connection, codigo, nuevo_nombre, nueva_direccion):
    try:
        cursor = connection.cursor()
        update_query = "UPDATE clientes SET  nombre=%s, direccion=%s WHERE codigo=%s"
        data = (nuevo_nombre, nueva_direccion, codigo)
        cursor.execute(update_query, data)
        connection.commit()
        print("Cliente actualizado con éxito.")
    except mysql.connector.Error as error:
        print("Error al actualizar cliente:", error)

def eliminar_clientes(connection, codigo):
    try:
        cursor = connection.cursor()
        delete_query = "DELETE FROM clientes WHERE codigo=%s"
        data = (codigo,)
        cursor.execute(delete_query, data)
        connection.commit()
        print("cliente eliminado con éxito.")
    except mysql.connector.Error as error:
        print("Error al eliminar cliente:", error)

def ventas_ayuda():
    print("Para poder utilizar el programa de Control de Ventas debe ingresar por la linea de comandos los datos separados por espacios, por ejemplo:")
    print("py proyecto.py --ventas-listar")
    print("py proyecto.py --ventas-crear codigo_venta codigo_producto codigo_cliente cantidad_producto")
    print("py proyecto.py --venta-anular codigo_venta")

#CONTROL DE VENTAS 
# Función para listar ventas desde la base de datos MySQL
def listar_ventas(connection):
    try:
        cursor = connection.cursor()
        cursor.execute("SELECT codigo_venta, codigo_producto, codigo_cliente, cantidad_productos, total_venta FROM ventas")
        ventas = cursor.fetchall()
        for venta in ventas:
            print("Código de Venta: {}, Código de Producto: {}, Código de Cliente: {}, Cantidad de productos: {}, Total de Venta: {}".format(*venta))
    except mysql.connector.Error as error:
        print("Error al listar ventas:", error)

def crear_venta(connection, codigo_venta, codigo_producto, codigo_cliente, cantidad_productos):
    try:
        # Verificar existencia disponible
        cursor = connection.cursor()
        cursor.execute("SELECT existencia, precio FROM inventario WHERE codigo=%s", (codigo_producto,))
        result = cursor.fetchone()

        if result is not None:
            existencia, precio = result

            cantidad_productos = int(cantidad_productos) 
            precio = float(precio) 

            if int(cantidad_productos) <= int(existencia):
                total_venta = cantidad_productos * precio

                update_query = "UPDATE inventario SET existencia = existencia - %s WHERE codigo = %s"
                cursor.execute(update_query, (cantidad_productos, codigo_producto))

                # Insertar la venta en la tabla de ventas
                insert_query = "INSERT INTO ventas (codigo_venta, codigo_producto, codigo_cliente, cantidad_productos, total_venta) VALUES (%s, %s, %s, %s, %s)"
                data = (codigo_venta, codigo_producto, codigo_cliente, cantidad_productos, total_venta)
                cursor.execute(insert_query, data)
                connection.commit()
                print("La venta se creó con éxito.")
            else:
                print("Lo sentimos, en este momento no hay suficiente existencia para realizar la venta.")
        else:
            print("Producto no encontrado en el inventario.")
    except mysql.connector.Error as error:
        print("Error al crear venta:", error)

# Función para anular una venta en la base de datos MySQL
def anular_venta(connection, codigo_venta):
    try:
        cursor = connection.cursor()
        cursor.execute("SELECT codigo_producto, cantidad_productos FROM ventas WHERE codigo_venta=%s", (codigo_venta,))
        venta = cursor.fetchone()

        if venta is not None:
            codigo_producto, cantidad_productos = venta
            # Actualizar existencia en la tabla de productos
            update_query = "UPDATE inventario SET existencia = existencia + %s WHERE codigo = %s"
            cursor.execute(update_query, (cantidad_productos, codigo_producto))

            # Eliminar la venta de la tabla de ventas
            delete_query = "DELETE FROM ventas WHERE codigo_venta = %s"
            cursor.execute(delete_query, (codigo_venta,))
            connection.commit()
            print("La venta fue anulada con éxito.")
        else:
            print("La venta con el código especificado no existe.")
    except mysql.connector.Error as error:
        print("Error al anular venta:", error)

#REPORTES BÁSICOS

def generar_informe_ventas_cliente(connection, cliente_id, destinatario):
    try:
        cursor = connection.cursor()
        cursor.execute("SELECT ventas.codigo_venta, ventas.codigo_producto, inventario.nombre, ventas.cantidad_productos, inventario.precio, ventas.total_venta FROM ventas INNER JOIN inventario ON ventas.codigo_producto = inventario.codigo WHERE ventas.codigo_cliente=%s", (cliente_id,))
        ventas_cliente = cursor.fetchall()

        # Crear un documento de Word
        doc = Document()
        header = doc.add_heading('Informe de Ventas por Cliente', 0)

        header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 
        run = header.runs[0] 
        run.bold = True 
        run.italic = True 
        run.font.name = 'Forte'

        magenta_color = RGBColor(255, 0, 255)
        run.font.color.rgb = magenta_color 

        # Agregar el nombre del cliente
        cursor.execute("SELECT nombre FROM clientes WHERE codigo=%s", (cliente_id,))
        nombre_cliente = cursor.fetchone()
        if nombre_cliente:
            Titulo2 = doc.add_paragraph(f'Cliente: {nombre_cliente[0]}', style='Heading1')

            Titulo2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 
            run = Titulo2.runs[0] 
            run.bold = True 
            run.italic = True 

            purple_color = RGBColor(128, 0, 128)
            run.font.color.rgb = purple_color 

        # Agregar la tabla de ventas
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        table.autofit = True

        heading_cells = table.rows[0].cells
        heading_cells[0].text = 'Código de Venta'
        heading_cells[1].text = 'Código de Producto'
        heading_cells[2].text = 'Nombre del Producto'
        heading_cells[3].text = 'Cantidad de Productos'
        heading_cells[4].text = 'precio del producto'
        heading_cells[5].text = 'Total de venta'

        for venta in ventas_cliente:
            cells = table.add_row().cells
            cells[0].text = str(venta[0])
            cells[1].text = str(venta[1])
            cells[2].text = str(venta[2])
            cells[3].text = str(venta[3])
            cells[4].text = str(venta[4])
            cells[5].text = str(venta[5])

        # Guardar el informe de ventas por cliente en un archivo
        nombre_archivo = f'Informe_Ventas_Cliente_{cliente_id}.docx'
        doc.save(nombre_archivo)

        enviar_correo(destinatario, f'Informe de Ventas Cliente {cliente_id}', f' En el archivo adjunto encontrarás el informe de ventas por  el cliente {cliente_id}.', [nombre_archivo])
    except mysql.connector.Error as error:
        print("Error al generar informe de ventas por cliente:", error)

def enviar_correo(destinatario, asunto, cuerpo, archivos_adjuntos):
    smtp_server = 'smtp.gmail.com' 
    smtp_port = 587 
    usuario_smtp = os.getenv('USUARIO')
    contraseña_smtp = os.getenv('CONTRASENA')

    msg = MIMEMultipart()
    msg['From'] = usuario_smtp
    msg['To'] = destinatario
    msg['Subject'] = asunto

    # Agregar el cuerpo del mensaje
    msg.attach(MIMEText(cuerpo, 'plain', 'utf-8'))

    # Adjuntar los archivos al mensaje
    for archivo in archivos_adjuntos:
        with open(archivo, "rb") as adjunto:
            part = MIMEApplication(adjunto.read(), Name=archivo)
            part['Content-Disposition'] = f'attachment; filename="{archivo}"'
            msg.attach(part)

    try:
        server = smtplib.SMTP(smtp_server, smtp_port)
        server.starttls()
        server.login(usuario_smtp, contraseña_smtp)
        server.sendmail(usuario_smtp, destinatario, msg.as_string())
        server.quit()
        print(f'El correo fue enviado exitosamente a {destinatario}')
    except smtplib.SMTPException as e:
        print(f'Error al enviar el correo: {str(e)}')

def informe_ventas_producto(connection, producto_id, destinatario):
    try:
        cursor = connection.cursor()
        cursor.execute("SELECT ventas.codigo_venta, ventas.codigo_producto, ventas.cantidad_productos, ventas.total_venta FROM ventas WHERE ventas.codigo_producto=%s", (producto_id,))
        ventas_producto = cursor.fetchall()

        doc = Document()
        header = doc.add_heading('Informe de Ventas por Producto', 0)

        header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 
        run = header.runs[0] 
        run.bold = True 
        run.italic = True 
        run.font.name = 'Consolas'

        purple_color = RGBColor(128, 0, 128)
        run.font.color.rgb = purple_color  

        cursor.execute("SELECT nombre FROM inventario WHERE codigo=%s", (producto_id,))
        nombre_producto = cursor.fetchone()
        if nombre_producto:
            Nombre = doc.add_paragraph(f'Producto: {nombre_producto[0]}', style='Heading1',)

            Nombre.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT 
            run = Nombre.runs[0] 
            run.bold = True 
            run.italic = True 

            magenta_color = RGBColor(255, 0, 255)
            run.font.color.rgb = magenta_color 

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.autofit = True

        heading_cells = table.rows[0].cells
        heading_cells[0].text = 'Código de Venta'
        heading_cells[1].text = 'Código de Producto'
        heading_cells[2].text = 'Cantidad de Productos'
        heading_cells[3].text = 'Total de venta'

        for venta in ventas_producto:
            cells = table.add_row().cells
            cells[0].text = str(venta[0])
            cells[1].text = str(venta[1])
            cells[2].text = str(venta[2])
            cells[3].text = str(venta[3])

        # Guardar el informe de ventas por producto en un archivo
        nombre_archivo = f'Informe_Ventas_Producto_{producto_id}.docx'
        doc.save(nombre_archivo)

        enviar_correo(destinatario, f'Informe de Ventas Producto {producto_id}', f' En el archivo adjunto encontrarás el informe de ventas por  el Producto {producto_id}.', [nombre_archivo])
    except mysql.connector.Error as error:
        print("Error al generar informe de ventas por producto:", error)

def enviar_correo(destinatario, asunto, cuerpo, archivos_adjuntos):
    smtp_server = 'smtp.gmail.com'
    smtp_port = 587 
    usuario_smtp = os.getenv('USUARIO')
    contraseña_smtp = os.getenv('CONTRASENA')

    msg = MIMEMultipart()
    msg['From'] = usuario_smtp
    msg['To'] = destinatario
    msg['Subject'] = asunto

    # Agregar el cuerpo del mensaje
    msg.attach(MIMEText(cuerpo, 'plain', 'utf-8'))

    # Adjuntar los archivos al mensaje
    for archivo in archivos_adjuntos:
        with open(archivo, "rb") as adjunto:
            part = MIMEApplication(adjunto.read(), Name=archivo)
            part['Content-Disposition'] = f'attachment; filename="{archivo}"'
            msg.attach(part)

    try:
        server = smtplib.SMTP(smtp_server, smtp_port)
        server.starttls()
        server.login(usuario_smtp, contraseña_smtp)
        server.sendmail(usuario_smtp, destinatario, msg.as_string())
        server.quit()
        print(f'Correo enviado a {destinatario}')
    except smtplib.SMTPException as e:
        print(f'Error al enviar el correo: {str(e)}')

def menu_interactivo(connection):
    connection = connect_to_database()
    if connection is None:
        return

    while True:
        print("\n Menú Principal:")
        print("1. Control de Inventario")
        print("2. Control de Clientes")
        print("3. Control de Ventas")
        print("4. Reportes básicos")
        print("5. Salir")

        choice = input("Elija una opción (1/2/3/4/5): ")

        if choice == "1":
            control_inventario_menu(connection)
        elif choice == "2":
            control_clientes_menu(connection)
        elif choice == "3":
            control_ventas_menu(connection)
        elif choice == "4":
            reportes_basicos_menu(connection)
        elif choice == "5":
            print ("Saliendo del programa... vuelva pronto")
            break
        else:
            print("Opción no válida. Intente de nuevo.")


def control_inventario_menu(connection):
    while True:
        print("\n Menú de Control de Inventario:")
        print("1. Listar productos")
        print("2. Crear producto")
        print("3. Actualizar producto")
        print("4. Editar existencia de producto")
        print("5. Eliminar producto")
        print("6. Volver al Menú Principal")

        choice = input("Elija una opción (1/2/3/4/5/6): \n")

        if choice == "1":
            listar_productos(connection)
        elif choice == "2":
            codigo =  input("Ingrese el codigo del producto: ")
            nombre =  input("Ingrese el nombre del producto: ")
            existencia = input("Ingrese la existencia del producto: ")
            proveedor = input("Ingrese el proveedor del producto: ")
            precio = input("Ingrese el precio del producto: ")
            crear_producto(connection, codigo, nombre, existencia, proveedor, precio)
            pass
        elif choice == "3":
            codigo = input("Ingrese el codigo del producto: ")
            nuevo_nombre = input("Ingrese el nuevo nombre dle producto: ")
            nueva_existencia = input("Ingrese la nueva existencia del producto: ")
            nuevo_proveedor = input("Ingrese el nuevo proveedor: ")
            nuevo_precio = input("Ingrese el nuevo precio del producto: ")
            actualizar_producto(connection, codigo, nuevo_nombre, nueva_existencia, nuevo_proveedor, nuevo_precio)
            pass
        elif choice == "4":
            codigo = input("Ingrese el codigo del producto: ")
            nueva_existencia = input("Ingrese la nueva existencia: ")
            editar_existencia(connection, codigo, nueva_existencia)
            pass
        elif choice == "5":
            codigo = input("Ingrese el codigo del producto que desea eliminar: ")
            eliminar_producto(connection, codigo)
            pass
        elif choice == "6":
            print("Volviendo al Menú Principal.")
            break
        else:
            print("Opción no válida. Intente de nuevo.")

def control_clientes_menu(connection):
    while True:
        print("\n Menú de Control de Clientes:")
        print("1. Listar clientes")
        print("2. Crear cliente")
        print("3. Actualizar cliente")
        print("4. Eliminar cliente")
        print("5. Volver al Menú Principal")

        choice = input("Elija una opción (1/2/3/4/5): \n")

        if choice == "1":
            listar_clientes(connection)
        elif choice == "2":
            codigo =  input("Ingrese el codigo del cliente: ")
            nombre =  input("Ingrese el nombre del cliente: ")
            direccion = input("Ingrese la direccion del cliente: ")
            crear_clientes(connection, codigo, nombre, direccion)
            pass
        elif choice == "3":
            codigo =  input("Ingrese el codigo del cliente: ")
            nuevo_nombre =  input("Ingrese el nuevo nombre del cliente: ")
            nueva_direccion = input("Ingrese la nueva direccion del cliente: ")
            actualizar_clientes(connection, codigo, nuevo_nombre, nueva_direccion)
            pass
        elif choice == "4":
            codigo = input("Ingrese el codigo del cliente que desea eliminar")
            eliminar_clientes(connection, codigo)
            pass
        elif choice == "5":
            print("Volviendo al Menú Principal.")
            break
        else:
            print("Opción no válida. Intente de nuevo.")

def control_ventas_menu(connection):
    while True:
        print("\n Menú de Control de Ventas:")
        print("1. Listar ventas")
        print("2. Crear venta")
        print("3. Anular venta")
        print("4. Volver al Menú Principal")

        choice = input("Elija una opción (1/2/3/4): \n")

        if choice == "1":
            listar_ventas(connection)
        elif choice == "2":
            codigo_venta = input("Ingrese el codigo de venta: ")
            codigo_producto = input("Ingrese el codigo del producto que desea vender: ")
            codigo_cliente = input("Ingrese el codigo del cliente: ")
            cantidad_productos = input("Ingrese la cantidad de productos que se venderán: ")
            crear_venta(connection, codigo_venta, codigo_producto, codigo_cliente, cantidad_productos)
            pass
        elif choice == "3":
            codigo_venta = input("Ingrese el codigo de la venta que desea eliminar: ")
            anular_venta(connection, codigo_venta)
            pass
        elif choice == "4":
            print("Volviendo al Menú Principal.")
            break
        else:
            print("Opción no válida. Intente de nuevo.")

def reportes_basicos_menu(connection):
    while True:
        print("\n Reportes básicos")
        print("1. Generar informe ventas por cliente")
        print("2. Generar informe ventas por producto")
        print("3. Volver al Menú principal")

        choice = input("Seleccione la opción (1/2/3): \n")

        if choice =="1":
            cliente_id = input("Ingrese el codigo del cliente: ")
            destinatario = input("Ingrese el correo electrónico del destinatario: ")
            generar_informe_ventas_cliente(connection, cliente_id, destinatario)
        elif choice == "2":
            producto_id = input("Ingrese el codigo del producto: ")
            destinatario = input("Ingrese el correo electrónico del destinatario: ")
            informe_ventas_producto(connection, producto_id, destinatario)
        elif choice == "3":
            print("Volviendo al Menú Principal.")
            break
        else:
            print("Opción no válida. Intente de nuevo.")

def principal():
    parser = argparse.ArgumentParser(description="Gestión de productos en la base de datos MySQL")
    parser.add_argument("--ayuda", action="store_true", help="Mostrar la ayuda para usar todo el programa")
    parser.add_argument("--inventario-ayuda", action="store_true", help="Mostrar la ayuda del inventario")
    parser.add_argument("--listar", action="store_true", help="Listar productos")
    parser.add_argument("--crear", nargs=5, metavar=("codigo", "nombre", "existencia", "proveedor", "precio"), help="Crear un nuevo producto")
    parser.add_argument("--actualizar", nargs=5, metavar=("codigo", "nuevo_nombre", "nueva_existencia", "nuevo_proveedor", "nuevo_precio"), help="Actualizar un producto")
    parser.add_argument("--editar-existencia", nargs=2, metavar=("codigo", "nueva_existencia"), help="Editar existencia de un producto")
    parser.add_argument("--eliminar", type=str, metavar="codigo", help="Eliminar un producto")
    parser.add_argument("--clientes-ayuda", action="store_true", help="Mostrar la ayuda de los clientes")
    parser.add_argument("--clientes-listar", action="store_true", help="Listar clientes")
    parser.add_argument("--clientes-crear", nargs=3, metavar=("codigo", "nombre", "direccion"), help="Crear un nuevo cliente")
    parser.add_argument("--clientes-actualizar", nargs=3, metavar=("codigo","nuevo_nombre","nueva_direccion"), help="Actualizar cliente")
    parser.add_argument("--clientes-eliminar", type=str, metavar="codigo", help="Eliminar cliente" )
    parser.add_argument("--ventas-ayuda", action="store_true", help="Mostrar la ayuda de las ventas")
    parser.add_argument("--ventas-listar",action="store_true", help="Listar ventas")
    parser.add_argument("--ventas-crear", nargs=4, metavar=("codigo_venta", "codigo_producto", "codigo_cliente","cantidad_productos"), help="Crear venta")
    parser.add_argument("--venta-anular", type=int, metavar="codigo_venta", help="Anular una venta por su código")
    parser.add_argument("--generar-informe-ventas-cliente", nargs=2, metavar=("cliente_id", "destinatario"), help="Generar informe de ventas por cliente y enviar por correo")
    parser.add_argument("--informe-ventas-producto", nargs=2, metavar=("producto_id", "destinatario"), help="Generar informe de ventas por producto y enviar por correo")
    parser.add_argument("--menu-interactivo", action="store_true", help="Mostrar el menú inventario")
    parser.add_argument("--salir", action="store_true", help="Salir del programa")
    
    args = parser.parse_args()

    # Conectar a la base de datos
    connection = connect_to_database()

    if connection:
        if args.ayuda:
            ayuda()
        elif args.inventario_ayuda:
            inventario_ayuda()
        elif args.listar:
            listar_productos(connection)
        elif args.crear:
            crear_producto(connection, *args.crear)
        elif args.actualizar:
            actualizar_producto(connection, *args.actualizar)
        elif args.editar_existencia:
            editar_existencia(connection, *args.editar_existencia)
        elif args.eliminar:
            eliminar_producto(connection, args.eliminar)
        elif args.clientes_ayuda:
            clientes_ayuda()
        elif args.clientes_listar:
            listar_clientes(connection)
        elif args.clientes_crear:
            crear_clientes(connection, *args.clientes_crear)
        elif args.clientes_actualizar:
            actualizar_clientes(connection, *args.clientes_actualizar)
        elif args.clientes_eliminar:
            eliminar_clientes(connection, args.clientes_eliminar)
        elif args.ventas_ayuda:
            ventas_ayuda()
        elif args.ventas_listar:
            listar_ventas(connection)
        elif args.ventas_crear:
            crear_venta(connection, *args.ventas_crear)
        elif args.venta_anular:
            anular_venta(connection, args.venta_anular)
        elif args.generar_informe_ventas_cliente:
            cliente_id, destinatario = args.generar_informe_ventas_cliente
            generar_informe_ventas_cliente(connection, cliente_id, destinatario)
        elif args.informe_ventas_producto:
            producto_id, destinatario = args.informe_ventas_producto
            informe_ventas_producto(connection, producto_id, destinatario)
        elif args.menu_interactivo:
            menu_interactivo(connection)
            connection.close()
        elif args.salir:
            print("Saliendo del programa... vuelva pronto.")
            parser.exit(0)
    else:
        print("No se pudo conectar a la base de datos.")

if __name__ == "__main__":
    principal()
